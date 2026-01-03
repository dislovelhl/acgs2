"""
Unit tests for DOCX document generator in the compliance-docs-service.

Tests cover:
- DOCX generator class initialization
- Custom styles for compliance documents
- Table builder functionality
- Report generation for all four compliance frameworks (SOC 2, ISO 27001, GDPR, EU AI Act)
- Buffer generation for streaming responses
- Error handling for unsupported frameworks
- Edge cases (minimal data, empty data)
- Output path handling
"""

import io
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

from generators.docx_generator import (
    ComplianceDOCXGenerator,
    ComplianceDOCXStyles,
    DOCXTableBuilder,
    generate_docx,
    generate_docx_to_buffer,
    _get_output_path,
    _ensure_output_dir,
)
from models.base import ComplianceFramework


class TestComplianceDOCXStyles:
    """Tests for ComplianceDOCXStyles class."""

    def test_styles_initialization(self):
        """Test that styles are initialized correctly."""
        doc = Document()
        styles = ComplianceDOCXStyles(doc)
        assert styles is not None
        assert styles.document is doc

    def test_color_definitions(self):
        """Test that color definitions are present."""
        doc = Document()
        styles = ComplianceDOCXStyles(doc)
        assert styles.PRIMARY_COLOR is not None
        assert styles.SECONDARY_COLOR is not None
        assert styles.SUCCESS_COLOR is not None
        assert styles.DANGER_COLOR is not None
        assert styles.WARNING_COLOR is not None

    def test_title_style_modified(self):
        """Test that Title style is properly modified."""
        doc = Document()
        ComplianceDOCXStyles(doc)
        title_style = doc.styles["Title"]
        assert title_style.font.size is not None
        assert title_style.font.bold is True

    def test_heading_styles_modified(self):
        """Test that heading styles are properly modified."""
        doc = Document()
        ComplianceDOCXStyles(doc)
        h1 = doc.styles["Heading 1"]
        h2 = doc.styles["Heading 2"]
        h3 = doc.styles["Heading 3"]
        assert h1.font.bold is True
        assert h2.font.bold is True
        assert h3.font.bold is True

    def test_custom_styles_created(self):
        """Test that custom styles are created."""
        doc = Document()
        ComplianceDOCXStyles(doc)
        # Check that custom styles were created
        style_names = [s.name for s in doc.styles]
        assert "ConfidentialNotice" in style_names
        assert "Metadata" in style_names
        assert "TableHeader" in style_names
        assert "TableCell" in style_names


class TestDOCXTableBuilder:
    """Tests for DOCXTableBuilder class."""

    def test_table_builder_initialization(self):
        """Test table builder initialization."""
        doc = Document()
        styles = ComplianceDOCXStyles(doc)
        builder = DOCXTableBuilder(doc, styles)
        assert builder is not None
        assert builder.document is doc
        assert builder.styles is styles

    def test_create_simple_table(self):
        """Test creating a simple table."""
        doc = Document()
        styles = ComplianceDOCXStyles(doc)
        builder = DOCXTableBuilder(doc, styles)
        headers = ["Col1", "Col2", "Col3"]
        rows = [["A", "B", "C"], ["D", "E", "F"]]
        table = builder.create_simple_table(headers, rows)
        assert table is not None
        # Header row + 2 data rows
        assert len(table.rows) == 3

    def test_create_simple_table_with_col_widths(self):
        """Test creating table with custom column widths."""
        doc = Document()
        styles = ComplianceDOCXStyles(doc)
        builder = DOCXTableBuilder(doc, styles)
        headers = ["ID", "Name"]
        rows = [["1", "Test"]]
        table = builder.create_simple_table(headers, rows, col_widths=[1.0, 3.0])
        assert table is not None

    def test_create_evidence_table(self):
        """Test creating an evidence table."""
        doc = Document()
        styles = ComplianceDOCXStyles(doc)
        builder = DOCXTableBuilder(doc, styles)
        records = [
            {
                "control_id": "CC1.1",
                "description": "Test evidence",
                "evidence_type": "document",
                "collected_at": datetime(2024, 6, 15, tzinfo=timezone.utc),
                "status": "compliant",
            }
        ]
        table = builder.create_evidence_table(records)
        assert table is not None

    def test_create_evidence_table_without_status(self):
        """Test creating evidence table without status column."""
        doc = Document()
        styles = ComplianceDOCXStyles(doc)
        builder = DOCXTableBuilder(doc, styles)
        records = [{"control_id": "A.5.1", "description": "Evidence"}]
        table = builder.create_evidence_table(records, include_status=False)
        assert table is not None
        # Without status, should have 4 columns instead of 5
        assert len(table.rows[0].cells) == 4

    def test_create_control_mapping_table(self):
        """Test creating a control mapping table."""
        doc = Document()
        styles = ComplianceDOCXStyles(doc)
        builder = DOCXTableBuilder(doc, styles)
        mappings = [
            {
                "soc2_control_id": "CC1.1",
                "guardrail_control_name": "GR-001",
                "mapping_rationale": "Test mapping",
                "coverage_percentage": 100,
            }
        ]
        table = builder.create_control_mapping_table(mappings, "soc2")
        assert table is not None

    def test_format_date_with_datetime(self):
        """Test date formatting with datetime object."""
        doc = Document()
        styles = ComplianceDOCXStyles(doc)
        builder = DOCXTableBuilder(doc, styles)
        dt = datetime(2024, 6, 15, tzinfo=timezone.utc)
        result = builder._format_date(dt)
        assert result == "2024-06-15"

    def test_format_date_with_none(self):
        """Test date formatting with None value."""
        doc = Document()
        styles = ComplianceDOCXStyles(doc)
        builder = DOCXTableBuilder(doc, styles)
        result = builder._format_date(None)
        assert result == "N/A"

    def test_format_status(self):
        """Test status formatting."""
        doc = Document()
        styles = ComplianceDOCXStyles(doc)
        builder = DOCXTableBuilder(doc, styles)
        assert builder._format_status("in_progress") == "In Progress"
        assert builder._format_status("") == "N/A"
        assert builder._format_status(None) == "N/A"

    def test_set_cell_shading(self):
        """Test setting cell shading."""
        doc = Document()
        styles = ComplianceDOCXStyles(doc)
        builder = DOCXTableBuilder(doc, styles)
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        # Should not raise an error
        builder._set_cell_shading(cell, "FF0000")

    def test_set_table_borders(self):
        """Test setting table borders."""
        doc = Document()
        styles = ComplianceDOCXStyles(doc)
        builder = DOCXTableBuilder(doc, styles)
        table = doc.add_table(rows=2, cols=2)
        # Should not raise an error
        builder._set_table_borders(table)


class TestComplianceDOCXGenerator:
    """Tests for ComplianceDOCXGenerator class."""

    def test_generator_initialization(self):
        """Test DOCX generator initialization."""
        generator = ComplianceDOCXGenerator()
        assert generator is not None
        assert generator.orientation == "portrait"

    def test_generator_initialization_landscape(self):
        """Test DOCX generator with landscape orientation."""
        generator = ComplianceDOCXGenerator(orientation="landscape")
        assert generator.orientation == "landscape"

    def test_create_document(self):
        """Test creating a new document."""
        generator = ComplianceDOCXGenerator()
        doc = generator._create_document()
        assert doc is not None
        assert generator._styles is not None
        assert generator._table_builder is not None

    def test_create_document_landscape(self):
        """Test creating document with landscape orientation."""
        generator = ComplianceDOCXGenerator(orientation="landscape")
        doc = generator._create_document()
        assert doc is not None
        # Check that orientation was set (width should be greater than height)
        section = doc.sections[0]
        assert section.page_width > section.page_height

    def test_add_section(self):
        """Test adding a section to the document."""
        generator = ComplianceDOCXGenerator()
        generator._create_document()
        initial_paragraphs = len(generator._document.paragraphs)
        generator._add_section("Test Section")
        # Should have added a heading
        assert len(generator._document.paragraphs) > initial_paragraphs

    def test_add_section_with_content(self):
        """Test adding a section with content."""
        generator = ComplianceDOCXGenerator()
        generator._create_document()
        generator._add_section("Test Section", content="Test content")
        paragraphs = [p.text for p in generator._document.paragraphs]
        assert "Test content" in paragraphs

    def test_add_paragraph(self):
        """Test adding a paragraph to the document."""
        generator = ComplianceDOCXGenerator()
        generator._create_document()
        generator._add_paragraph("Test paragraph content")
        paragraphs = [p.text for p in generator._document.paragraphs]
        assert "Test paragraph content" in paragraphs

    def test_add_paragraph_bold(self):
        """Test adding a bold paragraph."""
        generator = ComplianceDOCXGenerator()
        generator._create_document()
        generator._add_paragraph("Bold text", bold=True)
        assert len(generator._document.paragraphs) > 0

    def test_add_bullet_list(self):
        """Test adding a bullet list to the document."""
        generator = ComplianceDOCXGenerator()
        generator._create_document()
        items = ["Item 1", "Item 2", "Item 3"]
        generator._add_bullet_list(items)
        paragraphs = [p.text for p in generator._document.paragraphs]
        for item in items:
            assert item in paragraphs

    def test_add_numbered_list(self):
        """Test adding a numbered list to the document."""
        generator = ComplianceDOCXGenerator()
        generator._create_document()
        items = ["Step 1", "Step 2", "Step 3"]
        generator._add_numbered_list(items)
        paragraphs = [p.text for p in generator._document.paragraphs]
        for item in items:
            assert item in paragraphs

    def test_add_page_break(self):
        """Test adding a page break to the document."""
        generator = ComplianceDOCXGenerator()
        generator._create_document()
        # Should not raise an error
        generator._add_page_break()


class TestDOCXGeneratorSOC2:
    """Tests for SOC 2 DOCX report generation."""

    def test_generate_soc2_report(self, tmp_path, sample_soc2_report_data):
        """Test generating a SOC 2 DOCX report."""
        output_path = tmp_path / "soc2_test.docx"
        generator = ComplianceDOCXGenerator()
        result = generator.generate_soc2_report(sample_soc2_report_data, output_path)
        assert result == output_path
        assert output_path.exists()
        assert output_path.stat().st_size > 0

    def test_generate_soc2_report_opens_correctly(self, tmp_path, sample_soc2_report_data):
        """Test that generated SOC 2 DOCX can be opened."""
        output_path = tmp_path / "soc2_test.docx"
        generator = ComplianceDOCXGenerator()
        generator.generate_soc2_report(sample_soc2_report_data, output_path)
        # Try to open the document
        doc = Document(str(output_path))
        assert doc is not None
        assert len(doc.paragraphs) > 0

    def test_generate_soc2_report_default_path(self, sample_soc2_report_data):
        """Test generating SOC 2 report with default output path."""
        generator = ComplianceDOCXGenerator()
        result = generator.generate_soc2_report(sample_soc2_report_data)
        assert result.exists()
        assert "soc2_report" in result.name
        # Clean up
        result.unlink(missing_ok=True)

    def test_generate_soc2_report_with_system_description(self, tmp_path):
        """Test SOC 2 report with system description section."""
        data = {
            "organization_name": "Test Org",
            "audit_period_start": datetime(2024, 1, 1, tzinfo=timezone.utc),
            "audit_period_end": datetime(2024, 12, 31, tzinfo=timezone.utc),
            "system_description": {
                "system_description": "Test system description",
                "principal_service_commitments": ["Commitment 1", "Commitment 2"],
                "components": ["Component A", "Component B"],
            },
        }
        output_path = tmp_path / "soc2_sys_desc.docx"
        generator = ComplianceDOCXGenerator()
        result = generator.generate_soc2_report(data, output_path)
        assert result.exists()

    def test_generate_soc2_report_with_criteria_sections(self, tmp_path):
        """Test SOC 2 report with trust service criteria sections."""
        data = {
            "organization_name": "Test Org",
            "criteria_sections": [
                {
                    "criteria": "security",
                    "controls": [
                        {"control_id": "CC1.1", "title": "Test", "control_objective": "Obj"},
                    ],
                },
            ],
        }
        output_path = tmp_path / "soc2_criteria.docx"
        generator = ComplianceDOCXGenerator()
        result = generator.generate_soc2_report(data, output_path)
        assert result.exists()


class TestDOCXGeneratorISO27001:
    """Tests for ISO 27001 DOCX report generation."""

    def test_generate_iso27001_report(self, tmp_path, sample_iso27001_report_data):
        """Test generating an ISO 27001 DOCX report."""
        output_path = tmp_path / "iso27001_test.docx"
        generator = ComplianceDOCXGenerator()
        result = generator.generate_iso27001_report(sample_iso27001_report_data, output_path)
        assert result == output_path
        assert output_path.exists()
        assert output_path.stat().st_size > 0

    def test_generate_iso27001_report_opens_correctly(self, tmp_path, sample_iso27001_report_data):
        """Test that generated ISO 27001 DOCX can be opened."""
        output_path = tmp_path / "iso27001_test.docx"
        generator = ComplianceDOCXGenerator()
        generator.generate_iso27001_report(sample_iso27001_report_data, output_path)
        doc = Document(str(output_path))
        assert doc is not None

    def test_generate_iso27001_report_with_soa(self, tmp_path):
        """Test ISO 27001 report with Statement of Applicability."""
        data = {
            "organization_name": "Test Org",
            "isms_scope": "Test ISMS Scope",
            "statement_of_applicability": {
                "entries": [
                    {
                        "control_id": "A.5.1",
                        "control_title": "Security Policies",
                        "applicability": "applicable",
                        "implementation_status": "implemented",
                    },
                ],
            },
        }
        output_path = tmp_path / "iso27001_soa.docx"
        generator = ComplianceDOCXGenerator()
        result = generator.generate_iso27001_report(data, output_path)
        assert result.exists()

    def test_generate_iso27001_report_with_themes(self, tmp_path):
        """Test ISO 27001 report with theme sections and controls."""
        data = {
            "organization_name": "Test Org",
            "theme_sections": [
                {
                    "theme": "organizational",
                    "implementation_percentage": 85.5,
                    "controls": [
                        {"control_id": "A.5.1", "title": "Policy", "status": "implemented"},
                    ],
                },
            ],
        }
        output_path = tmp_path / "iso27001_themes.docx"
        generator = ComplianceDOCXGenerator()
        result = generator.generate_iso27001_report(data, output_path)
        assert result.exists()


class TestDOCXGeneratorGDPR:
    """Tests for GDPR DOCX report generation."""

    def test_generate_gdpr_report(self, tmp_path, sample_gdpr_report_data):
        """Test generating a GDPR DOCX report."""
        output_path = tmp_path / "gdpr_test.docx"
        generator = ComplianceDOCXGenerator()
        result = generator.generate_gdpr_report(sample_gdpr_report_data, output_path)
        assert result == output_path
        assert output_path.exists()
        assert output_path.stat().st_size > 0

    def test_generate_gdpr_report_opens_correctly(self, tmp_path, sample_gdpr_report_data):
        """Test that generated GDPR DOCX can be opened."""
        output_path = tmp_path / "gdpr_test.docx"
        generator = ComplianceDOCXGenerator()
        generator.generate_gdpr_report(sample_gdpr_report_data, output_path)
        doc = Document(str(output_path))
        assert doc is not None

    def test_generate_gdpr_report_with_controller_record(self, tmp_path):
        """Test GDPR report with controller record section."""
        data = {
            "organization_name": "Test Org",
            "entity_role": "controller",
            "controller_record": {
                "controller_name": "Test Controller",
                "controller_contact": {"name": "John Doe", "email": "john@test.com"},
                "dpo": {"name": "Jane Doe", "email": "jane@test.com"},
                "processing_activities": [
                    {"name": "Activity 1", "purposes": ["Purpose 1"], "status": "active"},
                ],
            },
        }
        output_path = tmp_path / "gdpr_controller.docx"
        generator = ComplianceDOCXGenerator()
        result = generator.generate_gdpr_report(data, output_path)
        assert result.exists()

    def test_generate_gdpr_report_with_data_flows(self, tmp_path):
        """Test GDPR report with data flows section."""
        data = {
            "organization_name": "Test Org",
            "data_flows": [
                {
                    "name": "Flow 1",
                    "data_source": "Source A",
                    "data_destination": "Dest B",
                    "crosses_border": True,
                },
            ],
        }
        output_path = tmp_path / "gdpr_flows.docx"
        generator = ComplianceDOCXGenerator()
        result = generator.generate_gdpr_report(data, output_path)
        assert result.exists()

    def test_generate_gdpr_report_with_security_measures(self, tmp_path):
        """Test GDPR report with security measures section."""
        data = {
            "organization_name": "Test Org",
            "security_measures": [
                {"description": "Encryption at rest"},
                {"description": "Access controls"},
            ],
        }
        output_path = tmp_path / "gdpr_security.docx"
        generator = ComplianceDOCXGenerator()
        result = generator.generate_gdpr_report(data, output_path)
        assert result.exists()


class TestDOCXGeneratorEUAIAct:
    """Tests for EU AI Act DOCX report generation."""

    def test_generate_euaiact_report(self, tmp_path, sample_euaiact_report_data):
        """Test generating an EU AI Act DOCX report."""
        output_path = tmp_path / "euaiact_test.docx"
        generator = ComplianceDOCXGenerator()
        result = generator.generate_euaiact_report(sample_euaiact_report_data, output_path)
        assert result == output_path
        assert output_path.exists()
        assert output_path.stat().st_size > 0

    def test_generate_euaiact_report_opens_correctly(self, tmp_path, sample_euaiact_report_data):
        """Test that generated EU AI Act DOCX can be opened."""
        output_path = tmp_path / "euaiact_test.docx"
        generator = ComplianceDOCXGenerator()
        generator.generate_euaiact_report(sample_euaiact_report_data, output_path)
        doc = Document(str(output_path))
        assert doc is not None

    def test_generate_euaiact_report_with_risk_assessments(self, tmp_path):
        """Test EU AI Act report with risk assessments."""
        data = {
            "organization_name": "Test Org",
            "ai_systems": [{"system_id": "AIS-001"}],
            "high_risk_systems_count": 1,
            "limited_risk_systems_count": 2,
            "minimal_risk_systems_count": 3,
            "risk_assessments": [
                {
                    "system_name": "AI System 1",
                    "risk_level": "high_risk",
                    "high_risk_category": "Category A",
                    "assessment_date": datetime(2024, 6, 15, tzinfo=timezone.utc),
                },
            ],
        }
        output_path = tmp_path / "euaiact_risk.docx"
        generator = ComplianceDOCXGenerator()
        result = generator.generate_euaiact_report(data, output_path)
        assert result.exists()

    def test_generate_euaiact_report_with_conformity_assessments(self, tmp_path):
        """Test EU AI Act report with conformity assessments."""
        data = {
            "organization_name": "Test Org",
            "ai_systems": [],
            "conformity_assessments": [
                {
                    "system_id": "AIS-001",
                    "assessment_type": "internal_control",
                    "assessment_result": "passed",
                    "certificate_number": "CERT-001",
                },
            ],
        }
        output_path = tmp_path / "euaiact_conformity.docx"
        generator = ComplianceDOCXGenerator()
        result = generator.generate_euaiact_report(data, output_path)
        assert result.exists()

    def test_generate_euaiact_report_with_tech_documentation(self, tmp_path):
        """Test EU AI Act report with technical documentation."""
        data = {
            "organization_name": "Test Org",
            "ai_systems": [],
            "technical_documentation": [
                {
                    "system_name": "AI System 1",
                    "documentation_status": "complete",
                },
            ],
        }
        output_path = tmp_path / "euaiact_tech_doc.docx"
        generator = ComplianceDOCXGenerator()
        result = generator.generate_euaiact_report(data, output_path)
        assert result.exists()

    def test_generate_euaiact_report_with_qms(self, tmp_path):
        """Test EU AI Act report with Quality Management System."""
        data = {
            "organization_name": "Test Org",
            "ai_systems": [],
            "quality_management_system": {
                "description": "Test QMS description",
                "policies": ["Policy 1", "Policy 2"],
            },
        }
        output_path = tmp_path / "euaiact_qms.docx"
        generator = ComplianceDOCXGenerator()
        result = generator.generate_euaiact_report(data, output_path)
        assert result.exists()


class TestGenerateDOCXFunction:
    """Tests for the generate_docx() main entry point function."""

    def test_generate_docx_soc2(self, tmp_path, sample_soc2_report_data):
        """Test generate_docx with SOC 2 framework."""
        output_path = tmp_path / "soc2.docx"
        result = generate_docx(sample_soc2_report_data, "soc2", output_path)
        assert result == output_path
        assert output_path.exists()

    def test_generate_docx_iso27001(self, tmp_path, sample_iso27001_report_data):
        """Test generate_docx with ISO 27001 framework."""
        output_path = tmp_path / "iso27001.docx"
        result = generate_docx(sample_iso27001_report_data, "iso27001", output_path)
        assert result == output_path
        assert output_path.exists()

    def test_generate_docx_gdpr(self, tmp_path, sample_gdpr_report_data):
        """Test generate_docx with GDPR framework."""
        output_path = tmp_path / "gdpr.docx"
        result = generate_docx(sample_gdpr_report_data, "gdpr", output_path)
        assert result == output_path
        assert output_path.exists()

    def test_generate_docx_euaiact(self, tmp_path, sample_euaiact_report_data):
        """Test generate_docx with EU AI Act framework."""
        output_path = tmp_path / "euaiact.docx"
        result = generate_docx(sample_euaiact_report_data, "euaiact", output_path)
        assert result == output_path
        assert output_path.exists()

    def test_generate_docx_with_framework_enum(self, tmp_path, sample_soc2_report_data):
        """Test generate_docx with ComplianceFramework enum."""
        output_path = tmp_path / "soc2_enum.docx"
        result = generate_docx(sample_soc2_report_data, ComplianceFramework.SOC2, output_path)
        assert result == output_path
        assert output_path.exists()

    def test_generate_docx_unsupported_framework(self, sample_soc2_report_data):
        """Test generate_docx raises error for unsupported framework."""
        with pytest.raises(ValueError) as exc_info:
            generate_docx(sample_soc2_report_data, "unsupported")
        assert "Unsupported framework" in str(exc_info.value)

    def test_generate_docx_case_insensitive_framework(self, tmp_path, sample_soc2_report_data):
        """Test generate_docx is case insensitive for framework."""
        output_path = tmp_path / "soc2_upper.docx"
        result = generate_docx(sample_soc2_report_data, "SOC2", output_path)
        assert result.exists()

    def test_generate_docx_landscape_orientation(self, tmp_path, sample_soc2_report_data):
        """Test generate_docx with landscape orientation."""
        output_path = tmp_path / "soc2_landscape.docx"
        result = generate_docx(
            sample_soc2_report_data, "soc2", output_path, orientation="landscape"
        )
        assert result.exists()


class TestGenerateDOCXToBuffer:
    """Tests for generate_docx_to_buffer() function."""

    def test_generate_docx_to_buffer_soc2(self, sample_soc2_report_data):
        """Test generating SOC 2 DOCX to buffer."""
        buffer = generate_docx_to_buffer(sample_soc2_report_data, "soc2")
        assert isinstance(buffer, io.BytesIO)
        content = buffer.getvalue()
        assert len(content) > 0
        # DOCX files are ZIP files, check for PK magic bytes
        assert content[:2] == b"PK"

    def test_generate_docx_to_buffer_iso27001(self, sample_iso27001_report_data):
        """Test generating ISO 27001 DOCX to buffer."""
        buffer = generate_docx_to_buffer(sample_iso27001_report_data, "iso27001")
        assert isinstance(buffer, io.BytesIO)
        assert buffer.getvalue()[:2] == b"PK"

    def test_generate_docx_to_buffer_gdpr(self, sample_gdpr_report_data):
        """Test generating GDPR DOCX to buffer."""
        buffer = generate_docx_to_buffer(sample_gdpr_report_data, "gdpr")
        assert isinstance(buffer, io.BytesIO)
        assert buffer.getvalue()[:2] == b"PK"

    def test_generate_docx_to_buffer_euaiact(self, sample_euaiact_report_data):
        """Test generating EU AI Act DOCX to buffer."""
        buffer = generate_docx_to_buffer(sample_euaiact_report_data, "euaiact")
        assert isinstance(buffer, io.BytesIO)
        assert buffer.getvalue()[:2] == b"PK"

    def test_generate_docx_to_buffer_unsupported_framework(self):
        """Test buffer generation raises error for unsupported framework."""
        with pytest.raises(ValueError) as exc_info:
            generate_docx_to_buffer({}, "unsupported")
        assert "Unsupported framework" in str(exc_info.value)

    def test_generate_docx_to_buffer_seek_position(self, sample_soc2_report_data):
        """Test that buffer seek position is at start after generation."""
        buffer = generate_docx_to_buffer(sample_soc2_report_data, "soc2")
        assert buffer.tell() == 0

    def test_generate_docx_to_buffer_can_be_opened(self, sample_soc2_report_data):
        """Test that buffered DOCX can be opened."""
        buffer = generate_docx_to_buffer(sample_soc2_report_data, "soc2")
        doc = Document(buffer)
        assert doc is not None


class TestOutputPathHandling:
    """Tests for output path handling functions."""

    def test_get_output_path_default(self):
        """Test getting default output path."""
        path = _get_output_path()
        assert path is not None
        assert "compliance-reports" in str(path)

    def test_get_output_path_from_env(self):
        """Test getting output path from environment variable."""
        with patch.dict("os.environ", {"COMPLIANCE_OUTPUT_PATH": "/custom/path"}):
            path = _get_output_path()
            assert str(path) == "/custom/path"

    def test_ensure_output_dir(self, tmp_path):
        """Test ensuring output directory exists."""
        custom_path = tmp_path / "compliance-test"
        with patch.dict("os.environ", {"COMPLIANCE_OUTPUT_PATH": str(custom_path)}):
            result = _ensure_output_dir()
            assert result.exists()
            assert result.is_dir()


class TestEdgeCases:
    """Tests for edge cases and error handling."""

    def test_generate_docx_empty_data(self, tmp_path):
        """Test generating DOCX with empty data."""
        output_path = tmp_path / "empty.docx"
        result = generate_docx({}, "soc2", output_path)
        assert result.exists()

    def test_generate_docx_minimal_data(self, tmp_path, minimal_docx_data):
        """Test generating DOCX with minimal data."""
        output_path = tmp_path / "minimal.docx"
        result = generate_docx(minimal_docx_data, "soc2", output_path)
        assert result.exists()

    def test_generate_docx_none_values(self, tmp_path):
        """Test generating DOCX with None values in data."""
        data = {
            "organization_name": None,
            "audit_period_start": None,
            "audit_period_end": None,
        }
        output_path = tmp_path / "none_values.docx"
        result = generate_docx(data, "soc2", output_path)
        assert result.exists()

    def test_generate_docx_special_characters(self, tmp_path):
        """Test generating DOCX with special characters in data."""
        data = {
            "organization_name": "Test & Company <Corp>",
            "report_title": 'Report "Test" 2024',
        }
        output_path = tmp_path / "special_chars.docx"
        result = generate_docx(data, "soc2", output_path)
        assert result.exists()

    def test_generate_docx_unicode_characters(self, tmp_path):
        """Test generating DOCX with unicode characters."""
        data = {
            "organization_name": "Test Org - Societe Francaise",
            "report_title": "Unicode Test Report",
        }
        output_path = tmp_path / "unicode.docx"
        result = generate_docx(data, "soc2", output_path)
        assert result.exists()

    def test_generate_docx_long_text(self, tmp_path):
        """Test generating DOCX with very long text content."""
        data = {
            "organization_name": "Test Org",
            "system_description": {
                "system_description": "A" * 10000,  # Very long description
            },
        }
        output_path = tmp_path / "long_text.docx"
        result = generate_docx(data, "soc2", output_path)
        assert result.exists()


class TestDateFormatting:
    """Tests for date formatting in DOCX generator."""

    def test_format_date_with_string(self):
        """Test formatting a date string."""
        generator = ComplianceDOCXGenerator()
        result = generator._format_date("2024-06-15")
        assert result == "2024-06-15"

    def test_format_date_with_datetime(self):
        """Test formatting a datetime object."""
        generator = ComplianceDOCXGenerator()
        dt = datetime(2024, 6, 15, tzinfo=timezone.utc)
        result = generator._format_date(dt)
        assert result == "2024-06-15"

    def test_format_date_with_none(self):
        """Test formatting None returns N/A."""
        generator = ComplianceDOCXGenerator()
        result = generator._format_date(None)
        assert result == "N/A"


# Fixtures for DOCX generator tests
@pytest.fixture
def sample_soc2_report_data():
    """Sample SOC 2 report data for DOCX testing."""
    return {
        "organization_name": "ACGS Test Corporation",
        "audit_period_start": datetime(2024, 1, 1, tzinfo=timezone.utc),
        "audit_period_end": datetime(2024, 12, 31, tzinfo=timezone.utc),
        "criteria_sections": [
            {
                "criteria": "security",
                "controls": [
                    {
                        "control_id": "CC1.1",
                        "title": "Commitment to Integrity",
                        "control_objective": "The entity demonstrates commitment.",
                    },
                ],
            },
        ],
        "control_mappings": [
            {
                "soc2_control_id": "CC1.1",
                "guardrail_control_id": "GR-001",
                "guardrail_control_name": "Ethics Guardrail",
                "mapping_rationale": "Maps to integrity commitment",
                "coverage_percentage": 100,
            },
        ],
    }


@pytest.fixture
def sample_iso27001_report_data():
    """Sample ISO 27001 report data for DOCX testing."""
    return {
        "organization_name": "ACGS Test Corporation",
        "isms_scope": "AI Guardrails Platform",
        "statement_of_applicability": {
            "entries": [
                {
                    "control_id": "A.5.1",
                    "control_title": "Information Security Policy",
                    "applicability": "applicable",
                    "implementation_status": "implemented",
                },
            ],
        },
        "theme_sections": [
            {"theme": "organizational", "implementation_percentage": 85.5},
        ],
    }


@pytest.fixture
def sample_gdpr_report_data():
    """Sample GDPR report data for DOCX testing."""
    return {
        "organization_name": "ACGS Test Corporation",
        "entity_role": "controller",
        "controller_record": {
            "controller_name": "ACGS Test Corporation",
            "controller_contact": {"name": "John Doe", "email": "john@test.com"},
            "dpo": {"name": "Jane Doe", "email": "dpo@test.com"},
            "processing_activities": [
                {
                    "name": "AI Model Training",
                    "purposes": ["Training", "Evaluation"],
                    "status": "active",
                },
            ],
        },
        "data_flows": [
            {
                "name": "User Data Flow",
                "data_source": "Application",
                "data_destination": "Cloud Storage",
                "crosses_border": False,
            },
        ],
    }


@pytest.fixture
def sample_euaiact_report_data():
    """Sample EU AI Act report data for DOCX testing."""
    return {
        "organization_name": "ACGS Test Corporation",
        "organization_role": "provider",
        "ai_systems": [
            {"system_id": "AIS-001", "system_name": "Guardrail Engine"},
        ],
        "high_risk_systems_count": 1,
        "limited_risk_systems_count": 2,
        "minimal_risk_systems_count": 0,
        "risk_assessments": [
            {
                "system_name": "Guardrail Engine",
                "risk_level": "high_risk",
                "high_risk_category": "Safety Component",
                "assessment_date": datetime(2024, 6, 15, tzinfo=timezone.utc),
            },
        ],
        "conformity_assessments": [
            {
                "system_id": "AIS-001",
                "assessment_type": "internal_control",
                "assessment_result": "passed",
                "certificate_number": "CERT-001",
            },
        ],
    }


@pytest.fixture
def minimal_docx_data():
    """Minimal data for DOCX testing."""
    return {"organization_name": "Test Org"}
