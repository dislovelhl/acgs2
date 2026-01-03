"""
DOCX Document Generator using python-docx
"""

import logging
from typing import Any, Dict

from docx import Document

logger = logging.getLogger(__name__)


class DOCXGenerator:
    """Generates DOCX compliance reports."""

    def generate(self, content: Dict[str, Any], output_path: str) -> str:
        """
        Generate DOCX report from content dictionary.

        Args:
            content: Dictionary containing report data
            output_path: Path to save the generated DOCX

        Returns:
            Path to the generated file
        """
        try:
            doc = Document()

            # Title
            if "title" in content:
                doc.add_heading(content["title"], 0)

            # Metadata
            if "metadata" in content:
                doc.add_heading("Report Metadata", level=1)
                table = doc.add_table(rows=1, cols=2)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Field"
                hdr_cells[1].text = "Value"

                for k, v in content["metadata"].items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(k)
                    row_cells[1].text = str(v)

                doc.add_page_break()

            # Sections
            if "sections" in content:
                for section in content["sections"]:
                    doc.add_heading(section.get("title", ""), level=1)
                    doc.add_paragraph(section.get("content", ""))

                    if "table" in section:
                        table_data = section["table"]
                        if table_data:
                            # Assume first row is header
                            rows = len(table_data)
                            cols = len(table_data[0]) if rows > 0 else 0

                            if rows > 0 and cols > 0:
                                table = doc.add_table(rows=rows, cols=cols)
                                table.style = "Table Grid"

                                for i, row_data in enumerate(table_data):
                                    row = table.rows[i]
                                    for j, cell_data in enumerate(row_data):
                                        row.cells[j].text = str(cell_data)

            doc.save(output_path)
            logger.info(f"DOCX generated successfully at {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Failed to generate DOCX: {e}")
            raise
