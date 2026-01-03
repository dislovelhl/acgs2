"""
DOCX document generator using python-docx
Constitutional Hash: cdd01ef066bc6cf2
"""

import logging
from pathlib import Path
from typing import Any, Dict

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .base import BaseGenerator

logger = logging.getLogger(__name__)


class DOCXGenerator(BaseGenerator):
    """DOCX document generator for compliance documents"""

    def __init__(self, output_dir: str = "/app/documents"):
        super().__init__(output_dir)
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available, DOCX generation will fail")

    def generate(self, data: Dict[str, Any], filename: str) -> Path:
        """
        Generate a DOCX document from data.

        Args:
            data: Document data dictionary
            filename: Output filename (without extension)

        Returns:
            Path to generated DOCX file
        """
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not available")

        output_path = self._get_output_path(filename, "docx")

        # Create document
        doc = Document()

        # Title
        title = data.get("title", "Compliance Document")
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add content based on document type
        doc_type = data.get("document_type", "generic")

        if doc_type == "risk_assessment":
            self._build_risk_assessment(doc, data)
        elif doc_type == "human_oversight":
            self._build_human_oversight(doc, data)
        elif doc_type == "compliance_checklist":
            self._build_compliance_checklist(doc, data)
        elif doc_type == "quarterly_report":
            self._build_quarterly_report(doc, data)
        else:
            self._build_generic(doc, data)

        # Save document
        doc.save(str(output_path))

        logger.info(f"Generated DOCX: {output_path}")
        return output_path

    def _build_risk_assessment(self, doc: Document, data: Dict[str, Any]):
        """Build risk assessment document content."""
        # Metadata
        doc.add_paragraph(f"System: {data.get('system_name', 'N/A')}")
        doc.add_paragraph(f"Assessment Date: {data.get('assessment_date', 'N/A')}")
        doc.add_paragraph(f"Assessor: {data.get('assessor_name', 'N/A')}")
        doc.add_paragraph()

        # Risk factors
        risk_factors = data.get("risk_factors", [])
        if risk_factors:
            doc.add_heading("Risk Factors", level=2)
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'

            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "Category"
            header_cells[1].text = "Description"
            header_cells[2].text = "Likelihood"
            header_cells[3].text = "Impact"
            header_cells[4].text = "Risk Level"

            # Data rows
            for factor in risk_factors:
                row_cells = table.add_row().cells
                row_cells[0].text = factor.get("category", "")
                row_cells[1].text = factor.get("description", "")
                row_cells[2].text = factor.get("likelihood", "")
                row_cells[3].text = factor.get("impact", "")
                row_cells[4].text = factor.get("risk_level", "")

            doc.add_paragraph()

        # Mitigation measures
        mitigations = data.get("mitigation_measures", [])
        if mitigations:
            doc.add_heading("Mitigation Measures", level=2)
            for i, measure in enumerate(mitigations, 1):
                doc.add_paragraph(f"{i}. {measure}", style='List Number')

    def _build_human_oversight(self, doc: Document, data: Dict[str, Any]):
        """Build human oversight document content."""
        # Metadata
        doc.add_paragraph(f"System: {data.get('system_name', 'N/A')}")
        doc.add_paragraph(f"Assessment Date: {data.get('assessment_date', 'N/A')}")
        doc.add_paragraph()

        # Oversight measures
        measures = data.get("oversight_measures", [])
        if measures:
            doc.add_heading("Human Oversight Measures", level=2)
            for measure in measures:
                doc.add_heading(measure.get("measure_type", "N/A"), level=3)
                doc.add_paragraph(f"Description: {measure.get('description', 'N/A')}")
                doc.add_paragraph(f"Responsible Role: {measure.get('responsible_role', 'N/A')}")
                doc.add_paragraph(f"Frequency: {measure.get('frequency', 'N/A')}")
                doc.add_paragraph()

    def _build_compliance_checklist(self, doc: Document, data: Dict[str, Any]):
        """Build compliance checklist document content."""
        # Metadata
        doc.add_paragraph(f"System: {data.get('system_name', 'N/A')}")
        doc.add_paragraph(f"Overall Status: {data.get('overall_status', 'N/A')}")
        doc.add_paragraph()

        # Findings
        findings = data.get("findings", [])
        if findings:
            doc.add_heading("Compliance Findings", level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'

            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "Article"
            header_cells[1].text = "Requirement"
            header_cells[2].text = "Status"
            header_cells[3].text = "Severity"

            # Data rows
            for finding in findings:
                row_cells = table.add_row().cells
                row_cells[0].text = finding.get("article", "")
                row_cells[1].text = finding.get("requirement", "")
                row_cells[2].text = finding.get("status", "")
                row_cells[3].text = finding.get("severity", "")

    def _build_quarterly_report(self, doc: Document, data: Dict[str, Any]):
        """Build quarterly report document content."""
        # Executive summary
        doc.add_heading("Executive Summary", level=2)
        doc.add_paragraph(data.get("executive_summary", "N/A"))
        doc.add_paragraph()

        # Metrics
        metrics = data.get("report_period", {})
        if metrics:
            doc.add_heading("Quarterly Metrics", level=2)
            doc.add_paragraph(f"Total Assessments: {metrics.get('total_assessments', 0)}")
            doc.add_paragraph(f"Compliant Systems: {metrics.get('compliant_systems', 0)}")
            doc.add_paragraph(f"Non-Compliant Systems: {metrics.get('non_compliant_systems', 0)}")

    def _build_generic(self, doc: Document, data: Dict[str, Any]):
        """Build generic document content."""
        content = data.get("content", "")
        if content:
            doc.add_paragraph(content)
