"""Constitutional Hash: cdd01ef066bc6cf2
DOCX Document Generator for Compliance Documentation Service

Generates professional compliance reports in DOCX format using python-docx.
Supports all compliance frameworks: SOC 2, ISO 27001, GDPR, and EU AI Act.

This module uses Document() for new files and provides structured document
generation with headings, paragraphs, tables, and professional styling.

Note: python-docx only supports .docx format (not legacy .doc format).
"""

import io
import logging
import os
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, BinaryIO, Dict, List, Optional, Union

try:
    from src.core.shared.types import DocumentData, JSONDict, JSONValue
except ImportError:
    from typing import Any, Dict, List, Union

    JSONPrimitive = Union[str, int, float, bool, None]
    JSONDict = Dict[str, Any]
    JSONList = List[Any]
    JSONValue = Union[JSONPrimitive, JSONDict, JSONList]
    DocumentData = Dict[str, JSONValue]

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table

from ..models.base import ComplianceFramework
from .base import BaseGenerator

logger = logging.getLogger(__name__)

# Default output path for generated DOCX files
_DEFAULT_OUTPUT_PATH = Path(tempfile.gettempdir()) / "compliance-reports"


def _get_output_path() -> Path:
    """
    Resolve output path from environment or use default.

    Returns:
        Path to output directory for generated DOCX files.
    """
    output_path = os.getenv("COMPLIANCE_OUTPUT_PATH")
    if output_path:
        return Path(output_path)
    return _DEFAULT_OUTPUT_PATH


def _ensure_output_dir() -> Path:
    """
    Ensure the output directory exists.

    Returns:
        Path to the output directory.
    """
    output_path = _get_output_path()
    output_path.mkdir(parents=True, exist_ok=True)
    return output_path


class ComplianceDOCXStyles:
    """
    Custom styles for compliance DOCX documents.

    Provides professional, consistent styling across all compliance reports.
    """

    # Color definitions
    PRIMARY_COLOR = RGBColor(0x1A, 0x36, 0x5D)  # Dark blue
    SECONDARY_COLOR = RGBColor(0x2C, 0x52, 0x82)  # Medium blue
    ACCENT_COLOR = RGBColor(0x42, 0x99, 0xE1)  # Light blue
    SUCCESS_COLOR = RGBColor(0x27, 0x67, 0x49)  # Green
    DANGER_COLOR = RGBColor(0xC5, 0x30, 0x30)  # Red
    WARNING_COLOR = RGBColor(0xC0, 0x56, 0x21)  # Orange
    TEXT_COLOR = RGBColor(0x2D, 0x37, 0x48)  # Dark gray
    MUTED_COLOR = RGBColor(0x4A, 0x55, 0x68)  # Gray

    def __init__(self, document: Document) -> None:
        """
        Initialize custom DOCX styles.

        Args:
            document: The Document instance to apply styles to.
        """
        self.document = document
        self._apply_custom_styles()

    def _apply_custom_styles(self) -> None:
        """Apply custom styles to the document."""
        styles = self.document.styles

        # Modify existing Title style
        title_style = styles["Title"]
        title_style.font.size = Pt(28)
        title_style.font.color.rgb = self.PRIMARY_COLOR
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(24)

        # Modify Heading 1
        h1_style = styles["Heading 1"]
        h1_style.font.size = Pt(18)
        h1_style.font.color.rgb = self.SECONDARY_COLOR
        h1_style.font.bold = True
        h1_style.paragraph_format.space_before = Pt(18)
        h1_style.paragraph_format.space_after = Pt(12)

        # Modify Heading 2
        h2_style = styles["Heading 2"]
        h2_style.font.size = Pt(14)
        h2_style.font.color.rgb = self.TEXT_COLOR
        h2_style.font.bold = True
        h2_style.paragraph_format.space_before = Pt(12)
        h2_style.paragraph_format.space_after = Pt(8)

        # Modify Heading 3
        h3_style = styles["Heading 3"]
        h3_style.font.size = Pt(12)
        h3_style.font.color.rgb = self.TEXT_COLOR
        h3_style.font.bold = True
        h3_style.paragraph_format.space_before = Pt(10)
        h3_style.paragraph_format.space_after = Pt(6)

        # Modify Normal style
        normal_style = styles["Normal"]
        normal_style.font.size = Pt(11)
        normal_style.font.color.rgb = self.TEXT_COLOR
        normal_style.paragraph_format.space_after = Pt(8)
        normal_style.paragraph_format.line_spacing = 1.15

        # Create custom styles if they don't exist
        self._create_custom_style(
            "ConfidentialNotice",
            WD_STYLE_TYPE.PARAGRAPH,
            font_size=12,
            font_color=self.DANGER_COLOR,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

        self._create_custom_style(
            "Metadata",
            WD_STYLE_TYPE.PARAGRAPH,
            font_size=10,
            font_color=self.MUTED_COLOR,
            alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        )

        self._create_custom_style(
            "TableHeader",
            WD_STYLE_TYPE.PARAGRAPH,
            font_size=10,
            font_color=RGBColor(0xFF, 0xFF, 0xFF),
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

        self._create_custom_style(
            "TableCell",
            WD_STYLE_TYPE.PARAGRAPH,
            font_size=9,
            font_color=self.TEXT_COLOR,
            alignment=WD_ALIGN_PARAGRAPH.LEFT,
        )

    def _create_custom_style(
        self,
        name: str,
        style_type: WD_STYLE_TYPE,
        font_size: int = 11,
        font_color: RGBColor = None,
        bold: bool = False,
        italic: bool = False,
        alignment: WD_ALIGN_PARAGRAPH = None,
    ) -> None:
        """
        Create a custom style if it doesn't exist.

        Args:
            name: Style name.
            style_type: Type of style (paragraph, character, etc.).
            font_size: Font size in points.
            font_color: Font color as RGBColor.
            bold: Whether text is bold.
            italic: Whether text is italic.
            alignment: Paragraph alignment.
        """
        styles = self.document.styles
        try:
            # Check if style already exists
            styles[name]
        except KeyError:
            # Create new style
            style = styles.add_style(name, style_type)
            style.font.size = Pt(font_size)
            if font_color:
                style.font.color.rgb = font_color
            style.font.bold = bold
            style.font.italic = italic
            if alignment and style_type == WD_STYLE_TYPE.PARAGRAPH:
                style.paragraph_format.alignment = alignment


class DOCXTableBuilder:
    """
    Builder class for creating formatted tables in DOCX documents.

    Provides methods for creating compliance-specific tables with
    consistent styling and formatting.
    """

    # Table style colors
    HEADER_BG_COLOR = "2C5282"  # Blue
    ALT_ROW_COLOR = "F7FAFC"  # Light gray
    BORDER_COLOR = "E2E8F0"  # Gray

    def __init__(self, document: Document, styles: ComplianceDOCXStyles) -> None:
        """
        Initialize the table builder.

        Args:
            document: The Document instance to add tables to.
            styles: The DOCX styles instance to use.
        """
        self.document = document
        self.styles = styles

    def create_simple_table(
        self,
        headers: List[str],
        rows: List[List[JSONValue]],
        col_widths: Optional[List[float]] = None,
    ) -> Table:
        """
        Create a simple table with headers and rows.

        Args:
            headers: List of column headers.
            rows: List of rows, each row is a list of cell values.
            col_widths: Optional list of column widths in inches.

        Returns:
            A formatted Table object.
        """
        # Create table with header row
        table = self.document.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Style header row
        header_cells = table.rows[0].cells
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = str(header_text)
            self._set_cell_shading(cell, self.HEADER_BG_COLOR)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)

        # Add data rows
        for row_idx, row_data in enumerate(rows):
            row_cells = table.add_row().cells
            for col_idx, cell_value in enumerate(row_data):
                cell = row_cells[col_idx]
                cell.text = str(cell_value) if cell_value is not None else "N/A"
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                run.font.size = Pt(9)
                run.font.color.rgb = self.styles.TEXT_COLOR

                # Apply alternating row colors
                if row_idx % 2 == 1:
                    self._set_cell_shading(cell, self.ALT_ROW_COLOR)

        # Set column widths if provided
        if col_widths and len(col_widths) == len(headers):
            for i, width in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = Inches(width)

        # Apply table borders
        self._set_table_borders(table)

        return table

    def create_evidence_table(
        self,
        evidence_records: List[JSONDict],
        include_status: bool = True,
    ) -> Table:
        """
        Create a table for evidence records.

        Args:
            evidence_records: List of evidence record dictionaries.
            include_status: Whether to include status column.

        Returns:
            A formatted Table for evidence display.
        """
        if include_status:
            headers = ["Control ID", "Evidence", "Type", "Collected", "Status"]
            col_widths = [1.0, 2.5, 1.0, 1.0, 1.0]
        else:
            headers = ["Control ID", "Evidence", "Type", "Collected"]
            col_widths = [1.0, 3.0, 1.0, 1.0]

        rows = []
        for record in evidence_records:
            row = [
                record.get("control_id", "N/A"),
                record.get("description", "N/A"),
                record.get("evidence_type", "N/A"),
                self._format_date(record.get("collected_at")),
            ]
            if include_status:
                row.append(self._format_status(record.get("status", "")))
            rows.append(row)

        return self.create_simple_table(headers, rows, col_widths)

    def create_control_mapping_table(
        self,
        mappings: List[JSONDict],
        framework: str,
    ) -> Table:
        """
        Create a table for control mappings.

        Args:
            mappings: List of control mapping dictionaries.
            framework: The compliance framework name.

        Returns:
            A formatted Table for control mappings.
        """
        headers = [
            f"{framework.upper()} Control",
            "Guardrail Control",
            "Mapping Rationale",
            "Coverage",
        ]
        col_widths = [1.2, 1.5, 3.0, 0.8]

        rows = []
        for mapping in mappings:
            coverage = mapping.get("coverage_percentage", mapping.get("coverage_level", 100))
            if isinstance(coverage, int):
                coverage_str = f"{coverage}%"
            else:
                coverage_str = str(coverage)

            row = [
                mapping.get("soc2_control_id")
                or mapping.get("iso27001_control_id")
                or mapping.get("control_id", "N/A"),
                mapping.get("guardrail_control_name", "N/A"),
                mapping.get("mapping_rationale", "N/A"),
                coverage_str,
            ]
            rows.append(row)

        return self.create_simple_table(headers, rows, col_widths)

    def _format_date(self, value: Optional[Union[datetime, str]]) -> str:
        """Format a datetime value for display."""
        if value is None:
            return "N/A"
        if isinstance(value, datetime):
            return value.strftime("%Y-%m-%d")
        return str(value)

    def _format_status(self, status: JSONValue) -> str:
        """Format a status value for display."""
        if not status:
            return "N/A"
        status_str = str(status).replace("_", " ").title()
        return status_str

    def _set_cell_shading(self, cell, color: str) -> None:
        """
        Set background shading for a table cell.

        Args:
            cell: The table cell to shade.
            color: Hex color code without # prefix.
        """
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), color)
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def _set_table_borders(self, table: Table) -> None:
        """
        Apply borders to a table.

        Args:
            table: The table to apply borders to.
        """
        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

        tbl_borders = OxmlElement("w:tblBorders")
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), self.BORDER_COLOR)
            tbl_borders.append(border)

        tbl_pr.append(tbl_borders)
        if tbl.tblPr is None:
            tbl.insert(0, tbl_pr)


class DOCXGenerator(BaseGenerator):
    """
    Main DOCX generator for compliance documentation.

    Generates professional compliance reports with consistent styling,
    proper structure, and support for all compliance frameworks.
    """

    def __init__(
        self,
        output_dir: Union[str, Path, None] = None,
        orientation: str = "portrait",
    ) -> None:
        """
        Initialize the DOCX generator.

        Args:
            output_dir: Directory for generated documents.
            orientation: Page orientation ('portrait' or 'landscape').
        """
        if output_dir is None:
            output_dir = _get_output_path()
        super().__init__(str(output_dir))
        self.orientation = orientation
        self._document: Optional[Document] = None
        self._styles: Optional[ComplianceDOCXStyles] = None
        self._table_builder: Optional[DOCXTableBuilder] = None

    def generate(
        self,
        data: DocumentData,
        output_path: Optional[Union[str, Path]] = None,
    ) -> Path:
        """
        Produce a compliance report using a unified interface.

        Args:
            data: The document data, including a mandatory "document_type" key.
            output_path: Where to save the resulting DOCX.

        Returns:
            The Path where the file was written.

        Raises:
            ValueError: If document_type is missing or unsupported.
        """
        doc_type = data.get("document_type")
        if not doc_type:
            raise ValueError("Direct generation requires 'document_type' in data.")

        # Handle output path resolution if provided
        if output_path:
            output_path = Path(output_path)
            # Ensure extension
            if not output_path.suffix:
                output_path = output_path.with_suffix(".docx")
            # Ensure generated in output_dir if not absolute
            if not output_path.is_absolute():
                output_path = self.output_dir / output_path

        if doc_type == "soc2_report":
            return self.generate_soc2_report(data, output_path)
        elif doc_type == "iso27001_report":
            return self.generate_iso27001_report(data, output_path)
        elif doc_type == "gdpr_report":
            return self.generate_gdpr_report(data, output_path)
        elif doc_type == "euaiact_report":
            return self.generate_euaiact_report(data, output_path)
        # Handle EU AI Act specific types from simpler generators
        elif doc_type in [
            "risk_assessment",
            "human_oversight",
            "compliance_checklist",
            "quarterly_report",
        ]:
            return self.generate_euaiact_report(data, output_path)
        else:
            # Default to EU AI Act report if unknown but it's the most common case for these
            return self.generate_euaiact_report(data, output_path)

    @property
    def document(self) -> Document:
        """Get the active Document instance."""
        from typing import cast

        return cast(Document, self._document)

    @property
    def styles(self) -> ComplianceDOCXStyles:
        """Get the active ComplianceDOCXStyles instance."""
        from typing import cast

        return cast(ComplianceDOCXStyles, self._styles)

    @property
    def table_builder(self) -> DOCXTableBuilder:
        """Get the active DOCXTableBuilder instance."""
        from typing import cast

        return cast(DOCXTableBuilder, self._table_builder)

    def _create_document(self) -> Document:
        """
        Create a new Document instance with custom styles.

        Returns:
            A new Document with styles applied.
        """
        self._document = Document()
        self._styles = ComplianceDOCXStyles(self._document)
        self._table_builder = DOCXTableBuilder(self._document, self._styles)

        # Set page orientation if landscape
        if self.orientation == "landscape":
            for section in self._document.sections:
                section.orientation = WD_ORIENT.LANDSCAPE
                # Swap width and height
                new_width = section.page_height
                new_height = section.page_width
                section.page_width = new_width
                section.page_height = new_height

        return self._document

    def _add_title_page(
        self,
        title: str,
        subtitle: Optional[str] = None,
        organization: Optional[str] = None,
        report_date: Optional[datetime] = None,
        confidentiality: str = "CONFIDENTIAL",
    ) -> None:
        """
        Add a title page to the document.

        Args:
            title: Main document title.
            subtitle: Optional subtitle.
            organization: Organization name.
            report_date: Report generation date.
            confidentiality: Confidentiality level.
        """
        assert self._document is not None
        assert self._styles is not None
        # Add blank paragraphs for spacing
        for _ in range(4):
            self.document.add_paragraph()

        # Title
        title_para = self.document.add_paragraph(title, style="Title")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        if subtitle:
            subtitle_para = self.document.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = subtitle_para.add_run(subtitle)
            run.font.size = Pt(16)
            run.font.color.rgb = self._styles.SECONDARY_COLOR

        # Add spacing
        for _ in range(2):
            self.document.add_paragraph()

        # Organization
        if organization:
            org_para = self.document.add_paragraph()
            org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = org_para.add_run(f"Prepared for: {organization}")
            run.font.size = Pt(14)
            run.font.color.rgb = self._styles.TEXT_COLOR

        # Report Date
        if report_date:
            date_para = self.document.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_str = report_date.strftime("%B %d, %Y")
            run = date_para.add_run(f"Report Date: {date_str}")
            run.font.size = Pt(12)
            run.font.color.rgb = self._styles.MUTED_COLOR

        # Add spacing before confidentiality notice
        for _ in range(4):
            self.document.add_paragraph()

        # Confidentiality Notice
        if confidentiality:
            conf_para = self.document.add_paragraph()
            conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = conf_para.add_run(confidentiality.upper())
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = self._styles.DANGER_COLOR

            notice_para = self.document.add_paragraph()
            notice_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            notice_run = notice_para.add_run(
                "This document contains confidential information. "
                "Distribution is restricted to authorized personnel only."
            )
            notice_run.font.size = Pt(10)
            notice_run.font.color.rgb = self._styles.MUTED_COLOR

        # Page break after title page
        self.document.add_page_break()

    def _add_section(
        self,
        title: str,
        content: Optional[str] = None,
        level: int = 1,
    ) -> None:
        """
        Add a section to the document.

        Args:
            title: Section title.
            content: Optional section content.
            level: Heading level (1, 2, or 3).
        """
        self.document.add_heading(title, level=level)

        if content:
            self.document.add_paragraph(content)

    def _add_paragraph(
        self,
        text: str,
        bold: bool = False,
        italic: bool = False,
    ) -> None:
        """
        Add a paragraph to the document.

        Args:
            text: Paragraph text.
            bold: Whether text is bold.
            italic: Whether text is italic.
        """
        para = self.document.add_paragraph()
        run = para.add_run(text)
        run.font.bold = bold
        run.font.italic = italic

    def _add_bullet_list(self, items: list[str]) -> None:
        """
        Add a bullet list to the document.

        Args:
            items: List of items to include.
        """
        for item in items:
            self.document.add_paragraph(item, style="List Bullet")

    def _add_numbered_list(self, items: list[str]) -> None:
        """
        Add a numbered list to the document.

        Args:
            items: List of items to include.
        """
        for item in items:
            self.document.add_paragraph(item, style="List Number")

    def _add_table(
        self,
        headers: List[str],
        rows: List[List[JSONValue]],
        col_widths: Optional[List[float]] = None,
    ) -> Table:
        """
        Add a table to the document.

        Args:
            headers: List of column headers.
            rows: List of rows.
            col_widths: Optional column widths in inches.

        Returns:
            The created Table object.
        """
        return self.table_builder.create_simple_table(headers, rows, col_widths)

    def _add_page_break(self) -> None:
        """Add a page break to the document."""
        self.document.add_page_break()

    def _save_document(
        self,
        output: Union[str, Path, BinaryIO],
    ) -> None:
        """
        Save the document to a file or buffer.

        Args:
            output: Output file path or file-like object.
        """
        if isinstance(output, (str, Path)):
            output_path = Path(output)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            self.document.save(str(output_path))
        else:
            self.document.save(output)

    def _format_date(self, value: Optional[Union[datetime, str]]) -> str:
        """Format a datetime value for display."""
        if value is None:
            return "N/A"
        if isinstance(value, datetime):
            return value.strftime("%Y-%m-%d")
        return str(value)

    def generate_soc2_report(
        self,
        report_data: DocumentData,
        output_path: Optional[Union[str, Path]] = None,
    ) -> Path:
        """
        Generate a SOC 2 Type II compliance report in DOCX format.

        Args:
            report_data: SOC 2 report data dictionary.
            output_path: Optional output file path.

        Returns:
            Path to the generated DOCX file.
        """
        self._create_document()

        org_name = report_data.get("organization_name", "Organization")
        audit_start = report_data.get("audit_period_start")
        audit_end = report_data.get("audit_period_end")

        # Title page
        self._add_title_page(
            title="SOC 2 Type II Report",
            subtitle="Service Organization Control Report",
            organization=org_name,
            report_date=datetime.now(timezone.utc),
        )

        # Executive Summary
        self._add_section("Executive Summary")
        self._add_paragraph(
            f"This SOC 2 Type II report covers the {org_name} system "
            f"for the audit period from {self._format_date(audit_start)} "
            f"to {self._format_date(audit_end)}."
        )

        # System Description
        system_desc = report_data.get("system_description", {})
        if system_desc:
            self._add_section("System Description")
            self._add_paragraph(
                system_desc.get("system_description", "System description not provided.")
            )

            if system_desc.get("principal_service_commitments"):
                self._add_section("Principal Service Commitments", level=2)
                self._add_bullet_list(system_desc["principal_service_commitments"])

            if system_desc.get("components"):
                self._add_section("System Components", level=2)
                self._add_bullet_list(system_desc["components"])

        # Trust Service Criteria
        criteria_sections = report_data.get("criteria_sections", [])
        if criteria_sections:
            self._add_page_break()
            self._add_section("Trust Service Criteria")

            for section in criteria_sections:
                criteria = section.get("criteria", "")
                criteria_name = str(criteria).replace("_", " ").title()
                self._add_section(f"{criteria_name} Controls", level=2)

                controls = section.get("controls", [])
                if controls:
                    rows = []
                    for ctrl in controls:
                        rows.append(
                            [
                                ctrl.get("control_id", "N/A"),
                                ctrl.get("title", "N/A"),
                                ctrl.get("control_objective", "N/A"),
                            ]
                        )

                    self._add_table(
                        headers=["Control ID", "Title", "Objective"],
                        rows=rows,
                        col_widths=[1.0, 2.0, 3.5],
                    )

        # Control Mappings
        mappings = report_data.get("control_mappings", [])
        if mappings:
            self._add_page_break()
            self._add_section("Guardrail Control Mappings")
            self._add_paragraph(
                "The following table shows how ACGS guardrail controls map to SOC 2 controls."
            )
            self.table_builder.create_control_mapping_table(
                [m if isinstance(m, dict) else m.model_dump() for m in mappings],
                "SOC2",
            )

        # Generate output path if not provided
        if output_path is None:
            output_dir = _ensure_output_dir()
            timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
            output_path = output_dir / f"soc2_report_{timestamp}.docx"

        self._save_document(output_path)

        logger.info(f"Generated SOC 2 DOCX report: {output_path}")
        return Path(output_path)

    def generate_iso27001_report(
        self,
        report_data: DocumentData,
        output_path: Optional[Union[str, Path]] = None,
    ) -> Path:
        """
        Generate an ISO 27001:2022 compliance report in DOCX format.

        Args:
            report_data: ISO 27001 report data dictionary.
            output_path: Optional output file path.

        Returns:
            Path to the generated DOCX file.
        """
        self._create_document()

        org_name = report_data.get("organization_name", "Organization")
        scope = report_data.get("isms_scope", "")

        # Title page
        self._add_title_page(
            title="ISO 27001:2022 Compliance Report",
            subtitle="Information Security Management System",
            organization=org_name,
            report_date=datetime.now(timezone.utc),
        )

        # ISMS Scope
        self._add_section("ISMS Scope")
        self._add_paragraph(scope or "ISMS scope not specified.")

        # Statement of Applicability
        soa = report_data.get("statement_of_applicability", {})
        if soa:
            self._add_section("Statement of Applicability")
            entries = soa.get("entries", [])

            if entries:
                rows = []
                for entry in entries[:20]:  # Limit for display
                    status = entry.get("implementation_status", "not_implemented")
                    status_display = str(status).replace("_", " ").title()
                    rows.append(
                        [
                            entry.get("control_id", "N/A"),
                            entry.get("control_title", "N/A"),
                            str(entry.get("applicability", "applicable")).replace("_", " ").title(),
                            status_display,
                        ]
                    )

                self._add_table(
                    headers=["Control ID", "Title", "Applicability", "Status"],
                    rows=rows,
                    col_widths=[1.0, 2.5, 1.2, 1.3],
                )

        # Theme Sections
        theme_sections = report_data.get("theme_sections", [])
        if theme_sections:
            self._add_page_break()
            self._add_section("Control Themes")

            for section in theme_sections:
                theme = section.get("theme", "")
                theme_name = str(theme).replace("_", " ").title()
                self._add_section(f"{theme_name} Controls", level=2)

                impl_pct = section.get("implementation_percentage", 0)
                self._add_paragraph(f"Implementation Progress: {impl_pct:.1f}%")

                controls = section.get("controls", [])
                if controls:
                    rows = []
                    for ctrl in controls[:10]:  # Limit per section
                        rows.append(
                            [
                                ctrl.get("control_id", "N/A"),
                                ctrl.get("title", "N/A"),
                                str(ctrl.get("status", "N/A")).replace("_", " ").title(),
                            ]
                        )

                    self._add_table(
                        headers=["Control ID", "Title", "Status"],
                        rows=rows,
                        col_widths=[1.2, 3.5, 1.3],
                    )

        # Generate output path if not provided
        if output_path is None:
            output_dir = _ensure_output_dir()
            timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
            output_path = output_dir / f"iso27001_report_{timestamp}.docx"

        self._save_document(output_path)

        logger.info(f"Generated ISO 27001 DOCX report: {output_path}")
        return Path(output_path)

    def generate_gdpr_report(
        self,
        report_data: DocumentData,
        output_path: Optional[Union[str, Path]] = None,
    ) -> Path:
        """
        Generate a GDPR Article 30 compliance report in DOCX format.

        Args:
            report_data: GDPR report data dictionary.
            output_path: Optional output file path.

        Returns:
            Path to the generated DOCX file.
        """
        self._create_document()

        org_name = report_data.get("organization_name", "Organization")
        entity_role = report_data.get("entity_role", "controller")

        # Title page
        self._add_title_page(
            title="GDPR Article 30 Report",
            subtitle="Records of Processing Activities",
            organization=org_name,
            report_date=datetime.now(timezone.utc),
        )

        # Entity Role
        self._add_section("Data Protection Role")
        self._add_paragraph(
            f"This organization operates as a data {str(entity_role).replace('_', ' ')} "
            f"under the General Data Protection Regulation (GDPR)."
        )

        # Controller Record
        controller_record = report_data.get("controller_record")
        if controller_record:
            self._add_section("Controller Information (Article 30(1))")

            contact = controller_record.get("controller_contact", {})
            if contact:
                self._add_paragraph(
                    f"Controller: {controller_record.get('controller_name', 'N/A')}"
                )
                self._add_paragraph(f"Contact: {contact.get('name', 'N/A')}")
                self._add_paragraph(f"Email: {contact.get('email', 'N/A')}")

            # DPO Information
            dpo = controller_record.get("dpo")
            if dpo:
                self._add_section("Data Protection Officer", level=2)
                self._add_paragraph(f"Name: {dpo.get('name', 'N/A')}")
                self._add_paragraph(f"Email: {dpo.get('email', 'N/A')}")

            # Processing Activities
            activities = controller_record.get("processing_activities", [])
            if activities:
                self._add_page_break()
                self._add_section("Processing Activities")

                rows = []
                for activity in activities[:15]:  # Limit for display
                    purposes = activity.get("purposes", [])
                    purposes_str = ", ".join(purposes[:2]) if purposes else "N/A"
                    if len(purposes) > 2:
                        # PERFORMANCE: Consider using list.append() + "".join() instead of += in loops
                        purposes_str += "..."

                    rows.append(
                        [
                            activity.get("name", "N/A"),
                            purposes_str,
                            activity.get("status", "N/A").replace("_", " ").title(),
                        ]
                    )

                self._add_table(
                    headers=["Activity", "Purposes", "Status"],
                    rows=rows,
                    col_widths=[2.0, 3.0, 1.5],
                )

        # Data Flows
        data_flows = report_data.get("data_flows", [])
        if data_flows:
            self._add_section("Data Flow Mappings")

            rows = []
            for flow in data_flows[:10]:  # Limit for display
                rows.append(
                    [
                        flow.get("name", "N/A"),
                        flow.get("data_source", "N/A"),
                        flow.get("data_destination", "N/A"),
                        "Yes" if flow.get("crosses_border") else "No",
                    ]
                )

            self._add_table(
                headers=["Flow Name", "Source", "Destination", "Cross-Border"],
                rows=rows,
                col_widths=[1.8, 1.5, 1.5, 1.2],
            )

        # Security Measures
        security_measures = report_data.get("security_measures", [])
        if security_measures:
            self._add_section("Technical and Organisational Measures")
            self._add_paragraph(
                "The following security measures are implemented to protect personal data:"
            )
            self._add_bullet_list([m.get("description", str(m)) for m in security_measures[:10]])

        # Generate output path if not provided
        if output_path is None:
            output_dir = _ensure_output_dir()
            timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
            output_path = output_dir / f"gdpr_report_{timestamp}.docx"

        self._save_document(output_path)

        logger.info(f"Generated GDPR DOCX report: {output_path}")
        return Path(output_path)

    def generate_euaiact_report(
        self,
        report_data: DocumentData,
        output_path: Optional[Union[str, Path]] = None,
    ) -> Path:
        """
        Generate an EU AI Act compliance report in DOCX format.

        Args:
            report_data: EU AI Act report data dictionary.
            output_path: Optional output file path.

        Returns:
            Path to the generated DOCX file.
        """
        self._create_document()

        org_name = report_data.get("organization_name", "Organization")
        org_role = report_data.get("organization_role", "provider")

        # Title page
        self._add_title_page(
            title="EU AI Act Compliance Report",
            subtitle="Regulation (EU) 2024/1689",
            organization=org_name,
            report_date=datetime.now(timezone.utc),
        )

        # Organization Role
        self._add_section("Organization Role")
        self._add_paragraph(
            f"This organization operates as an AI system {str(org_role).replace('_', ' ')} "
            f"under the EU AI Act."
        )

        # AI Systems Summary
        ai_systems = report_data.get("ai_systems", [])
        high_risk_count = report_data.get("high_risk_systems_count", 0)
        limited_risk_count = report_data.get("limited_risk_systems_count", 0)
        minimal_risk_count = report_data.get("minimal_risk_systems_count", 0)

        self._add_section("AI Systems Overview")
        overview_items = [
            f"Total AI Systems: {len(ai_systems)}",
            f"High-Risk Systems: {high_risk_count}",
            f"Limited-Risk Systems: {limited_risk_count}",
            f"Minimal-Risk Systems: {minimal_risk_count}",
        ]
        self._add_bullet_list(overview_items)

        # Risk Assessments
        risk_assessments = report_data.get("risk_assessments", [])
        if risk_assessments:
            self._add_page_break()
            self._add_section("Risk Assessments")

            rows = []
            for assessment in risk_assessments[:10]:  # Limit for display
                risk_level = assessment.get("risk_level", "N/A")
                rows.append(
                    [
                        assessment.get("system_name", "N/A"),
                        str(risk_level).replace("_", " ").title(),
                        assessment.get("high_risk_category", "N/A") or "N/A",
                        self._format_date(assessment.get("assessment_date")),
                    ]
                )

            self._add_table(
                headers=["System", "Risk Level", "Category", "Assessment Date"],
                rows=rows,
                col_widths=[2.0, 1.2, 1.5, 1.3],
            )

        # Conformity Assessments
        conformity_assessments = report_data.get("conformity_assessments", [])
        if conformity_assessments:
            self._add_section("Conformity Assessments")

            rows = []
            for assessment in conformity_assessments[:10]:  # Limit for display
                rows.append(
                    [
                        assessment.get("system_id", "N/A"),
                        str(assessment.get("assessment_type", "N/A")).replace("_", " ").title(),
                        assessment.get("assessment_result", "N/A").title(),
                        assessment.get("certificate_number", "N/A") or "N/A",
                    ]
                )

            self._add_table(
                headers=["System ID", "Assessment Type", "Result", "Certificate"],
                rows=rows,
                col_widths=[1.5, 1.5, 1.2, 1.8],
            )

        # Technical Documentation
        tech_docs = report_data.get("technical_documentation", [])
        if tech_docs:
            self._add_section("Technical Documentation (Annex IV)")
            self._add_paragraph("The following AI systems have complete technical documentation:")
            doc_items = []
            for doc in tech_docs[:10]:
                system_name = doc.get("system_name", "Unknown System")
                status = doc.get("documentation_status", "incomplete")
                doc_items.append(f"{system_name}: {status.replace('_', ' ').title()}")
            self._add_bullet_list(doc_items)

        # Quality Management System
        qms = report_data.get("quality_management_system")
        if qms:
            self._add_section("Quality Management System (Article 17)")
            self._add_paragraph(qms.get("description", "QMS description not provided."))

            policies = qms.get("policies", [])
            if policies:
                self._add_section("Documented Policies", level=2)
                self._add_bullet_list(policies[:10])

        # Generate output path if not provided
        if output_path is None:
            output_dir = _ensure_output_dir()
            timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
            output_path = output_dir / f"euaiact_report_{timestamp}.docx"

        self._save_document(output_path)

        logger.info(f"Generated EU AI Act DOCX report: {output_path}")
        return Path(output_path)


def generate_docx(
    report_data: DocumentData,
    framework: Union[str, ComplianceFramework],
    output_path: Optional[Union[str, Path]] = None,
    orientation: str = "portrait",
) -> Path:
    """
    Generate a DOCX compliance report for the specified framework.

    This is the main entry point for DOCX generation.

    Args:
        report_data: Report data dictionary containing all compliance information.
        framework: Compliance framework (soc2, iso27001, gdpr, euaiact).
        output_path: Optional output file path. If not provided, a default path is used.
        orientation: Page orientation ('portrait' or 'landscape').

    Returns:
        Path to the generated DOCX file.

    Raises:
        ValueError: If the framework is not supported.

    Example:
        >>> from acgs2.services.compliance_docs.src.generators.docx_generator import (
        ...     generate_docx
        ... )
        >>> path = generate_docx(
        ...     report_data={"organization_name": "Acme Corp", ...},
        ...     framework="soc2",
        ... )
        >>> print(f"Report generated at: {path}")
    """
    generator = DOCXGenerator(orientation=orientation)

    # Normalize framework
    if isinstance(framework, ComplianceFramework):
        framework_str = framework.value
    else:
        framework_str = str(framework).lower()

    if framework_str == "soc2":
        return generator.generate_soc2_report(report_data, output_path)
    elif framework_str == "iso27001":
        return generator.generate_iso27001_report(report_data, output_path)
    elif framework_str == "gdpr":
        return generator.generate_gdpr_report(report_data, output_path)
    elif framework_str == "euaiact":
        return generator.generate_euaiact_report(report_data, output_path)
    else:
        raise ValueError(
            f"Unsupported framework: {framework}. "
            f"Supported frameworks: soc2, iso27001, gdpr, euaiact"
        )


def generate_docx_to_buffer(
    report_data: DocumentData,
    framework: Union[str, ComplianceFramework],
    orientation: str = "portrait",
) -> io.BytesIO:
    """
    Generate a DOCX compliance report to an in-memory buffer.

    This is useful for streaming responses without writing to disk.

    Args:
        report_data: Report data dictionary containing all compliance information.
        framework: Compliance framework (soc2, iso27001, gdpr, euaiact).
        orientation: Page orientation ('portrait' or 'landscape').

    Returns:
        BytesIO buffer containing the generated DOCX.

    Raises:
        ValueError: If the framework is not supported.

    Example:
        >>> buffer = generate_docx_to_buffer(report_data, "soc2")
        >>> # Use buffer.getvalue() to get bytes for streaming
    """
    buffer = io.BytesIO()
    generator = DOCXGenerator(orientation=orientation)

    # Normalize framework
    if isinstance(framework, ComplianceFramework):
        framework_str = framework.value
    else:
        framework_str = str(framework).lower()

    # Create document and generate content
    generator._create_document()

    org_name = report_data.get("organization_name", "Organization")

    if framework_str == "soc2":
        generator._add_title_page(
            title="SOC 2 Type II Report",
            subtitle="Service Organization Control Report",
            organization=org_name,
            report_date=datetime.now(timezone.utc),
        )
        generator._add_section("Report Content")
        generator._add_paragraph(
            "Full SOC 2 report content. See generate_docx() for complete implementation."
        )
    elif framework_str == "iso27001":
        generator._add_title_page(
            title="ISO 27001:2022 Compliance Report",
            subtitle="Information Security Management System",
            organization=org_name,
            report_date=datetime.now(timezone.utc),
        )
        generator._add_section("Report Content")
        generator._add_paragraph(
            "Full ISO 27001 report content. See generate_docx() for complete implementation."
        )
    elif framework_str == "gdpr":
        generator._add_title_page(
            title="GDPR Article 30 Report",
            subtitle="Records of Processing Activities",
            organization=org_name,
            report_date=datetime.now(timezone.utc),
        )
        generator._add_section("Report Content")
        generator._add_paragraph(
            "Full GDPR report content. See generate_docx() for complete implementation."
        )
    elif framework_str == "euaiact":
        generator._add_title_page(
            title="EU AI Act Compliance Report",
            subtitle="Regulation (EU) 2024/1689",
            organization=org_name,
            report_date=datetime.now(timezone.utc),
        )
        generator._add_section("Report Content")
        generator._add_paragraph(
            "Full EU AI Act report content. See generate_docx() for complete implementation."
        )
    else:
        raise ValueError(
            f"Unsupported framework: {framework}. "
            f"Supported frameworks: soc2, iso27001, gdpr, euaiact"
        )

    generator._save_document(buffer)
    buffer.seek(0)
    return buffer
